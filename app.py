import streamlit as st
import pandas as pd
from io import BytesIO
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import os, tempfile
from docx2pdf import convert as _docx2pdf_convert
try:
    from docx2pdf import convert as _docx2pdf_convert
except Exception:
    _docx2pdf_convert = None

# -------------------- Data helpers --------------------
def make_course(name: str, code: str, cfu: int = 6, dept: str = "DIETI", year: str = "Second", semester: str = "Second"):
    # Normalize semester to words (first/second) and support 'first&second'
    s = str(semester).strip()
    mapping = {
        "I": "first", "1": "first", "First": "first", "first": "first",
        "II": "second", "2": "second", "Second": "second", "second": "second",
        "first&Second": "first&second", "First&Second": "first&second", "first&second": "first&second",
    }
    sem_norm = mapping.get(s, s)
    return {
        "name": name,
        "code": code,
        "cfu": int(cfu),
        "dept": dept,
        "year": year,
        "semester": sem_norm,
    }

def course_label(c):
    return f"{c['name']} ({c['code']}, {c['cfu']} CFU)"

# Fixed second-year components (now like normal courses)
FIXED_COMPONENTS = [
    make_course("ALTRE ATTIVITA", "12568", 6, "DIETI – LM Data Science", "Second", "second"),
    make_course("TESI DI LAUREA", "U2848", 16, "DIETI – LM Data Science", "Second", "second"),
    make_course("TIROCINIO/STAGE", "U4319", 8, "DIETI – LM Data Science", "Second", "second"),
]

# -------------------- Document helpers --------------------
def academic_year_to_aa_format(academic_year: str) -> str:
    """Convert '2025-2026' -> '2025/26'. If already like '2025/26', return as-is."""
    if "-" in academic_year:
        try:
            y1, y2 = academic_year.split("-")
            return f"{y1}/{str(int(y2)%100).zfill(2)}"
        except Exception:
            return academic_year
    return academic_year


def build_study_plan_docx(
    name: str,
    matricula: str,
    pob: str,
    dob_str: str,
    phone: str,
    email: str,
    academic_year: str,
    year_of_degree: str,
    degree_type: str,
    degree_name: str,
    main_path: str,
    sub_path: str,
    courses: list,
    watermark_text: str = None,
) -> BytesIO:
    doc = Document()

    # Optional header watermark for PSI
    if watermark_text:
        try:
            section = doc.sections[0]
            header = section.header
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            hr = hp.add_run(watermark_text)
            hr.font.size = Pt(40)
            hr.bold = True
            hr.font.color.rgb = RGBColor(200, 200, 200)
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # --- Base font setup (adjustable) ---
    try:
        normal = doc.styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
    except Exception:
        pass

    def add_p(text, *, align=None, bold=False, italic=False, size=12):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        # Ensure font for EastAsia too (Word quirk)
        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if align is not None:
            p.alignment = align
        return p

    def add_justified_with_bold(parts, *, size=12):
        """Add a justified paragraph built from (text, bold) tuples."""
        p = doc.add_paragraph()
        for text, is_bold in parts:
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            if r._element.rPr is not None and r._element.rPr.rFonts is not None:
                r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            r.font.size = Pt(size)
            r.bold = bool(is_bold)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # Header block (centered, styled)
    add_p("Università degli Studi di Napoli Federico II", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    add_p("Corso di Studio", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_p(f"Laurea Magistrale in {degree_name}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
    add_p("Piano di Studi", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
    add_p(f"A.A {academic_year_to_aa_format(academic_year)}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

    # Indirizzo (sub path)
    add_p(f"Indirizzo: {sub_path}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_p("Da consegnare al Coordinatore del Corso, Prof. Giuseppe Longo", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
    doc.add_paragraph("")

    # Body text (anagraphica)
    add_justified_with_bold([
        ("Il/La sottoscritto/a ", False),
        (name, True),
        (", matr. ", False),
        (matricula, True),
        (", nato/a a ", False),
        (pob, True),
        (" il ", False),
        (dob_str, True),
        (", cell. ", False),
        (phone, True),
        (", e-mail ", False),
        (email, True),
    ], size=12)

    add_justified_with_bold([
        ("iscritto/a nell’A.A. ", False),
        (academic_year_to_aa_format(academic_year), True),
        (" al ", False),
        (year_of_degree, True),
        (" anno del Corso di ", False),
        (degree_type, True),
        (" in ", False),
        (degree_name, True),
        (", chiede alla Commissione di Coordinamento Didattico del Corso di Studio l’approvazione del presente Piano di Studio (PdS).", False),
    ], size=12)

    doc.add_paragraph("")

    # Table (6 columns x 8 rows: 1 header + 7 items) with full borders
    table = doc.add_table(rows=8, cols=6)
    table.style = 'Table Grid'  # ensures visible borders on all cells
    hdr = table.rows[0].cells
    headers = [
        "Insegnamento",
        "Corso Di Laurea Da Cui È Offerto",
        "Codice Insegnamento",
        "CFU",
        "Anno",
        "Semestre",
    ]
    for i, h in enumerate(headers):
        # clear default content and add bold centered text
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fill rows with: 2 curricular + 2 free + 3 fixed = 7 rows
    for i, c in enumerate(courses[:7], start=1):
        row = table.rows[i].cells
        # Insegnamento
        row[0].text = c["name"]
        # Dipartimento / Corso di laurea offerente
        row[1].text = c["dept"]
        # Codice
        row[2].text = c["code"]
        # CFU (center)
        row[3].text = str(c["cfu"])
        for para in row[3].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Anno (center)
        row[4].text = str(c["year"])
        for para in row[4].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Semestre (center)
        row[5].text = str(c["semester"])
        for para in row[5].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Modalità di compilazione (title bold)
    add_p("Modalità di compilazione:", bold=True, size=12)
    bullets = [
        "Si possono includere nel PdS sia insegnamenti consigliati dal Corso di Studio (elencati e di immediata approvazione) sia insegnamenti offerti presso l’Ateneo (riportare nome insegnamento, codice esame, Corso di Studio) purchè costituiscano un percorso didattico complementare, coerente con il Corso di Studio",
        "É ammesso il superamento del numero dei CFU previsti",
    ]
    for b in bullets:
        p = doc.add_paragraph(b)
        p.style = 'List Bullet'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("")
    today_str = date.today().strftime("%d/%m/%Y")
    # Signature line: left = firma dello studente, right = Napoli (date)
    sig_table = doc.add_table(rows=1, cols=2)
    try:
        sig_table.style = 'Table Normal'
    except Exception:
        pass
    left_p = sig_table.cell(0, 0).paragraphs[0]
    lr = left_p.add_run("firma dello studente")
    lr.italic = True
    lr.font.name = 'Times New Roman'
    lr.font.size = Pt(12)
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    right_p = sig_table.cell(0, 1).paragraphs[0]
    rr = right_p.add_run(f"Napoli ({today_str})")
    rr.font.name = 'Times New Roman'
    rr.font.size = Pt(12)
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save to buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def convert_docx_to_pdf_bytes(docx_buf: BytesIO) -> BytesIO | None:
    """Convert a DOCX buffer to PDF using docx2pdf (MS Word/AppleScript backend).
    Returns a BytesIO with PDF bytes, or None if conversion failed or unavailable.
    """
    if _docx2pdf_convert is None:
        return None
    try:
        with tempfile.TemporaryDirectory() as td:
            in_path = os.path.join(td, "plan.docx")
            out_path = os.path.join(td, "plan.pdf")
            with open(in_path, "wb") as f:
                f.write(docx_buf.getvalue())
            _docx2pdf_convert(in_path, out_path)
            if not os.path.exists(out_path):
                return None
            with open(out_path, "rb") as f:
                pdf_bytes = f.read()
            return BytesIO(pdf_bytes)
    except Exception:
        return None


def main():
    st.set_page_config(page_title="Master's Study Plan", page_icon="🎓", layout="wide")
    st.markdown(
        """
        <style>
            .title {text-align: center; color: #4CAF50;}
            .sidebar .sidebar-content {background-color: #f0f2f6;}
            .stButton>button {background-color: #4CAF50; color: white; border-radius: 8px;}
            .stDownloadButton>button {background-color: #008CBA; color: white; border-radius: 8px;}
            .card {padding: 1rem; border-radius: 12px; background: #ffffff; box-shadow: 0 2px 10px rgba(0,0,0,0.06);}            
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.title("🎓 Master's Study Plan Generator")

    # -------------------- Auth --------------------
    teacher_logged_in = False
    with st.sidebar:
        st.subheader("🔑 Teacher Login")
        teacher_id = st.text_input("Enter ID:", type="password")
        teacher_pass = st.text_input("Enter Password:", type="password")
        if teacher_id == "Professor@unina.it" and teacher_pass == "ISPD03":
            teacher_logged_in = True
            st.success("✅ Logged in successfully!")

    # -------------------- Predefined catalog --------------------
    if "catalog" not in st.session_state:
        st.session_state.catalog = {
            "Curriculum FUNDAMENTAL SCIENCES": {
                "FSE/PH - CURRICULUM FUNDAMENTAL SCIENCES/PHYSICS INSPIRED METHODOLOGIES": [
                    make_course("Advanced Statistical Learning and Modeling", "U5450", 12, "DIETI – LM Data Science", "Second", "first"),
                    make_course("Physics Informed Machine Learning", "U****", 6, "DIETI - LM Data Science", "Second", "second"),
                ],
                "PDS FSE/MM - CURRICULUM FUNDAMENTAL SCIENCES/MATHEMATICAL METHODOLOGIES": [
                    make_course("Algorithms and Parallel Computing and Computational Complexity", "U5430", 12, "DMRC - LM Ing. Matematica D71", "Second", "first&second"),
                    make_course("Operational Research", "U1624", 6, "DMRC - LM Ing. Matematica D71", "Second", "first"),
                ],
            },
            "Curriculum INFORMATION TECHNOLOGIES": {
                "PDS ITE/TS - CURRICULUM INFORMATION TECHNOLOGIES/TEXT AND SPEECH PROCESSING": [
                    make_course("Multimedia Information Retrieval and Text Mining", "U5441", 12, "DIETI – LM Data Science", "Second", "first&second"),
                    make_course("Natural Language Processing", "U3539", 6, "DIETI – LM Informatica", "Second", "second"),
                ],
                "PDS ITE/SV - CURRICULUM INFORMATION TECHNOLOGIES/SIGNAL AND VIDEO PROCESSING": [
                    make_course("Information Theory and Signals Theory", "U5444", 12, "DMRC - LM Ing. Matematica D71", "Second", "first&second"),
                    make_course("Image and Video Processing for Autonomous Driving", "U3423", 6, "DII - LM Autonomous Vehicle Engineering", "Second", "second"),
                ],
                "PDS ITE/RS - CURRICULUM INFORMATION TECHNOLOGIES/ STATISTICS AND ROBOTICS FOR HEALTH": [
                    make_course("Advanced Statistical Learning and Modeling", "U5450", 12, "DMRC - DIETI – LM Data Science", "Second", "first"),
                    make_course("Robotics for Bioengineering", "U1579", 6, "LM Ing. Automazione e Robotica", "Second", "second"),
                ],
                "PDS ITE/IA - CURRICULUM INFORMATION TECHNOLOGIES/INDUSTRIAL APPLICATIONS": [
                    make_course("Advanced Statistical Learning and Modeling", "U5450", 12, "DMRC - DIETI – LM Data Science", "Second", "first"),
                    make_course("Statistical Methods for Industrial Process Monitoring", "U2659", 6, "DMRC - LM Ing. Matematica D71", "Second", "first"),
                ],
                "PDS ITE/AI - CURRICULUM INFORMATION TECHNOLOGIES/DATA SECURITY": [
                    make_course("Data Security and Computer Forensics", "U5447", 12, "DMRC - DIETI – LM Informatica", "Second", "second"),
                    make_course("Algorithm Design", "U3524", 6, "DIETI – LM Data Science", "Second", "first"),
                ],
            },
            "Curriculum PUBLIC ADMINISTRATION, ECONOMY AND MANAGEMENT – ECO": {
                "PDS ECO - CURRICULUM PUBLIC ADMINISTRATION, ECONOMY AND MANAGEMENT": [
                    make_course("Computational Statistical and Generalized Linear Models", "U5453", 12, "DIETI – LM Data Science", "Second", "first"),
                    make_course("Financial Time Series Analysis", "U6373", 6, "DISES – LM Economics and Finance DH5", "Second", "first"),
                ],
            },
            "Curriculum INTELLIGENT SYSTEMS - ISY": {
                "PDS ISY - CURRICULUM INTELLIGENT SYSTEMS": [
                    make_course("Computational Intelligence and Machine Learning for Physics", "U5460", 12, "DFEP – LM Physics and DIETI – LM Data Science", "Second", "second"),
                    make_course("Generative Artificial Intelligence", "U****", 6, "DIETI – LM Data Science", "Second", "first"),
                ],
            },
        }

    if "specializations" in st.session_state and isinstance(st.session_state["specializations"], dict):
        it = st.session_state.catalog.get("Curriculum INFORMATION TECHNOLOGIES", {})
        it.update(st.session_state.specializations)
        st.session_state.catalog["Curriculum INFORMATION TECHNOLOGIES"] = it
        del st.session_state["specializations"]

    if "free_choice_courses" not in st.session_state:
        st.session_state.free_choice_courses = [
            make_course("Advanced Statistical Learning and Modeling", "U5450", 12, "DIETI – LM Data Science", "Second", "I"),
            make_course("AI Systems Engineering", "U5494", 6, "DIETI – LM Ing. Informatica", "Second", "I"),
            make_course("Astroinformatics", "U1205", 6, "DFEP – LM Fisica", "Second", "II"),
            make_course("Biometric Systems", "U3525", 6, "DIETI – LM Informatica", "Second", "II"),
            make_course("Computational Intelligence", "U7219", 6, "DIETI – LM Data Science", "Second", "II"),
            make_course("Computational Statistical and Generalized Linear Models", "U5453", 12, "DIETI – LM Data Science", "Second", "I"),
            make_course("Computer Vision", "U3523", 6, "DIETI – LM Informatica", "Second", "I"),
            make_course("Data Security", "U2652", 6, "DIETI – LM Data Science", "Second", "I"),
            make_course("Data Visualization", "U2658", 6, "DIETI – LM Data Science", "Second", "II"),
            make_course("Generative Artificial Intelligence", "U7215", 6, "DIETI – LM Data Science", "Second", "I"),
            make_course("Financial Time Series Analysis", "U6373", 6, "DISES – LM Econ. and Finance", "Second", "I"),
            make_course("Human robot interaction", "U3536", 6, "DIETI – LM Informatica", "Second", "I"),
            make_course("Image and Video Processing for Autonomous Driving", "U3423", 6, "DII - LM Autonomous Vehicle Engineering", "Second", "II"),
            make_course("Information Systems and Business Intelligence", "U3546", 6, "DIETI – LM Ing. Informatica", "Second", "I"),
            make_course("Information Theory", "U1644", 6, "DMRC - LM Ing. Matematica", "Second", "I"),
            make_course("Methods for Artificial Intelligence", "U3522", 6, "DIETI – LM Informatica", "Second", "II"),
            make_course("Natural Language Processing", "U3539", 6, "DIETI – LM Informatica", "Second", "II"),
            make_course("Physics Informed Machine Learning", "NI", 6, "DIETI – LM Data Science", "Second", "II"),
            make_course("Preference learning", "U6641", 6, "DISES – LM Economia e Commercio", "Second", "I"),
            make_course("Reliability and Risk in Aerospace Engineering", "U3835", 6, "DII – LM Ing. Aerospaziale", "Second", "II"),
            make_course("Robotics Lab", "U2325", 6, "DIETI – LM Ing. Automazione e Robotica", "Second", "I"),
            make_course("Software Architecture Design", "U5937", 6, "DIETI – LM Ing. Informatica", "Second", "I"),
            make_course("Speech Processing", "U6636", 6, "DIETI – LM Data Science", "Second", "II"),
            make_course("Statistical Methods for Industrial Process Monitoring", "U2659", 6, "DMRC - LM Ing. Matematica", "Second", "I"),
            make_course("SW and methods for statistical analysis of economic data", "U6640", 6, "DIETI – LM Data Science", "Second", "II"),
            make_course("Techniques of Text Analysis and Computational Linguistic", "U6635", 6, "DIETI – LM Data Science", "Second", "I"),
            make_course("Text Mining", "U5902", 6, "DIETI – LM Data Science", "Second", "I"),
            make_course("Advanced Microeconomics", "25880", 12, "DISES – LM Economics and Finance", "Second", "I"),
            make_course("Advanced Macroeconomics", "25881", 12, "DISES – LM Economics and Finance", "Second", "II"),
            make_course("Economics of Regulation", "27381", 6, "DISES – LM Economics and Finance", "Second", "II"),
            make_course("Financial Econometrics", "27382", 6, "DISES – LM Economics and Finance", "Second", "II"),
            make_course("Mathematics for Economics and Finance", "25884", 12, "DISES – LM Economics and Finance", "Second", "I"),
        ]

    recommended_by_path = {
        "Curriculum INFORMATION TECHNOLOGIES": {
            "PDS ITE/TS - CURRICULUM INFORMATION TECHNOLOGIES/TEXT AND SPEECH PROCESSING": [
                "Techniques of Text Analysis and Computational Linguistic",
                "Speech Processing",
            ],
            "PDS ITE/SV - CURRICULUM INFORMATION TECHNOLOGIES/SIGNAL AND VIDEO PROCESSING": [
                "Biometric Systems",
                "Image and Video Processing for Autonomous Driving",
            ],
            "PDS ITE/RS - CURRICULUM INFORMATION TECHNOLOGIES/ STATISTICS AND ROBOTICS FOR HEALTH": [
                "Methods for Artificial Intelligence",
                "Human robot interaction",
            ],
            "PDS ITE/IA - CURRICULUM INFORMATION TECHNOLOGIES/INDUSTRIAL APPLICATIONS": [
                "Environmental risk monitoring and evaluation",
                "Statistical Methods for Industrial Process Monitoring",
            ],
            "PDS ITE/AI - CURRICULUM INFORMATION TECHNOLOGIES/DATA SECURITY": [
                "Information Systems and Business Intelligence",
                "Computer Vision",
            ],
        },
        "Curriculum FUNDAMENTAL SCIENCES": {
            "PDS FSE/MM - CURRICULUM FUNDAMENTAL SCIENCES/MATHEMATICAL METHODOLOGIES": [
                "x-Informatics",
                "AI and Quantum Computing",
                "Real and Functional Analysis",
                "Signal Theory",
                "Information Theory",
                "Probability Theory",
            ],
            "FSE/PH - CURRICULUM FUNDAMENTAL SCIENCES/PHYSICS INSPIRED METHODOLOGIES": [
                "Astroinformatics",
                "AI and Quantum Computing",
                "Computational Thermodynamics",
                "Quantum Computing Systems",
                "Mathematical Physics Models",
            ],
        },
        "Curriculum PUBLIC ADMINISTRATION, ECONOMY AND MANAGEMENT – ECO": {
            "PDS ECO - CURRICULUM PUBLIC ADMINISTRATION, ECONOMY AND MANAGEMENT": [
                "Mathematics for Economics and Finance",
                "SW and methods for statistical analysis of economic data",
                "Advanced Microeconomics",
                "Advanced Macroeconomics",
                "Economics of Regulation",
                "Financial Econometrics",
            ]
        },
        "Curriculum INTELLIGENT SYSTEMS - ISY": {
            "PDS ISY - CURRICULUM INTELLIGENT SYSTEMS": [
                "Astroinformatics",
                "Quantum Computing Systems",
                "Computational Intelligence",
                "Generative Artificial Intelligence",
            ]
        },
    }

    # -------------------- Catalog overview --------------------
    with st.expander("📚 Catalog Overview (Codes, CFUs, Dept, Year, Semester)"):
        rows = []
        for main_path, subpaths in st.session_state.catalog.items():
            for sub_path, courses in subpaths.items():
                for idx, c in enumerate(courses, start=1):
                    rows.append({
                        "Type": "Curricular",
                        "Main Path": main_path,
                        "Sub Path": sub_path,
                        "Slot": f"Curricular {idx}",
                        "Course": c["name"],
                        "Code": c["code"],
                        "CFU": c["cfu"],
                        "Dept": c["dept"],
                        "Year": c["year"],
                        "Semester": c["semester"],
                    })
        for c in st.session_state.free_choice_courses:
            rows.append({
                "Type": "Free Choice",
                "Main Path": "—",
                "Sub Path": "—",
                "Slot": "—",
                "Course": c["name"],
                "Code": c["code"],
                "CFU": c["cfu"],
                "Dept": c["dept"],
                "Year": c["year"],
                "Semester": c["semester"],
            })
        for c in FIXED_COMPONENTS:
            rows.append({
                "Type": "Fixed",
                "Main Path": "—",
                "Sub Path": "—",
                "Slot": "—",
                "Course": c["name"],
                "Code": c["code"],
                "CFU": c["cfu"],
                "Dept": c["dept"],
                "Year": c["year"],
                "Semester": c["semester"],
            })
        df = pd.DataFrame(rows)
        st.dataframe(df, use_container_width=True)

    # -------------------- Teacher tools --------------------
    if teacher_logged_in:
        with st.expander("👨‍🏫 Teacher: Add Main/Sub Path & Courses"):
            all_main_paths = list(st.session_state.catalog.keys()) + ["➕ Create new main path…"]
            choice = st.selectbox("Select Main Path or create new:", all_main_paths, index=0, placeholder="Type to search…")
            new_main = None
            if choice == "➕ Create new main path…":
                new_main = st.text_input("New Main Path Name")
            main_selected = new_main if new_main else choice

            sub_path = st.text_input("Sub Path Name")
            st.markdown("Defaults reflect program rules: Curricular I = 12 CFU, Curricular II = 6 CFU; Dept/Year/Sem default to DIETI/Second/Second")
            with st.form("add_spec_form"):
                col1, col2 = st.columns(2)
                with col1:
                    c1_name = st.text_input("Curricular Course I Name", key="c1n")
                    c1_code = st.text_input("Course I Code", value="U7XXX", key="c1c")
                    c1_cfu = st.number_input("Course I CFU", min_value=1, max_value=30, value=12, key="c1f")
                    c1_dept = st.text_input("Course I Department", value="DIETI", key="c1d")
                    c1_year = st.selectbox("Course I Year", ["First", "Second"], index=1, key="c1y")
                    c1_sem = st.selectbox("Course I Semester", ["First", "Second"], index=1, key="c1s")
                with col2:
                    c2_name = st.text_input("Curricular Course II Name", key="c2n")
                    c2_code = st.text_input("Course II Code", value="U7YYY", key="c2c")
                    c2_cfu = st.number_input("Course II CFU", min_value=1, max_value=30, value=6, key="c2f")
                    c2_dept = st.text_input("Course II Department", value="DIETI", key="c2d")
                    c2_year = st.selectbox("Course II Year", ["First", "Second"], index=1, key="c2y")
                    c2_sem = st.selectbox("Course II Semester", ["First", "Second"], index=1, key="c2s")
                submitted = st.form_submit_button("➕ Add / Update Sub Path")
            if submitted:
                if main_selected and sub_path and c1_name and c2_name and c1_code and c2_code:
                    if main_selected not in st.session_state.catalog:
                        st.session_state.catalog[main_selected] = {}
                    st.session_state.catalog[main_selected][sub_path] = [
                        make_course(c1_name, c1_code, c1_cfu, c1_dept, c1_year, c1_sem),
                        make_course(c2_name, c2_code, c2_cfu, c2_dept, c2_year, c2_sem),
                    ]
                    st.success(f"✅ Saved sub path '{sub_path}' under main path '{main_selected}'.")
                else:
                    st.error("⚠ Please fill all required fields (names & codes).")

            st.markdown("Add a free choice course (defaults: DIETI / Second / Second; recommended CFU = 6):")
            with st.form("add_free_form"):
                f_name = st.text_input("Free Choice Course Name")
                f_code = st.text_input("Course Code", value="U8ZZZ")
                f_cfu = st.number_input("CFU", min_value=1, max_value=30, value=6)
                f_dept = st.text_input("Department", value="DIETI")
                f_year = st.selectbox("Year", ["First", "Second"], index=1, key="fy")
                f_sem = st.selectbox("Semester", ["First", "Second"], index=1, key="fs")
                submitted_free = st.form_submit_button("➕ Add Free Choice Course")
            if submitted_free:
                if f_name and f_code:
                    if all(fc["name"] != f_name for fc in st.session_state.free_choice_courses):
                        st.session_state.free_choice_courses.append(make_course(f_name, f_code, f_cfu, f_dept, f_year, f_sem))
                        st.success(f"✅ Course '{f_name}' added!")
                    else:
                        st.warning("A free choice course with this name already exists.")
                else:
                    st.error("⚠ Please enter a course name and code.")

    # -------------------- Student workflow --------------------
    with st.expander("🎓 Student: Select Your Study Plan", expanded=True):
        # --- Student personal details ---
        st.markdown("#### 🧑‍🎓 Student Details")
        ca, cb, cc = st.columns(3)
        with ca:
            name = st.text_input("Name")
            pob = st.text_input("Place of Birth")
            phone = st.text_input("Phone Number")
        with cb:
            matricula = st.text_input("Matricula")
            dob = st.date_input(
                "Date of Birth",
                value=date(2000, 1, 1),  # default shown when the page loads
                min_value=date(1900, 1, 1),  # allow going well before 2015
                max_value=date.today(),  # no future DoB
                help="Select your birth date (you can navigate years).",
            )
            email = st.text_input("Institutional Email")
        with cc:
            # Academic year helper options
            today = date.today()
            start_year = today.year if today.month >= 7 else today.year - 1
            acad_options = [f"{start_year-1}-{start_year}", f"{start_year}-{start_year+1}", f"{start_year+1}-{start_year+2}"]
            academic_year = st.selectbox("Academic Year", acad_options, index=1)
            year_of_degree = st.selectbox("Year of Degree", ["First", "Second"], index=1)
            degree_type = st.text_input("Degree Type", value="LAUREA MAGISTRALE")
            degree_name = st.text_input("Degree Name", value="DATA SCIENCE")

        st.markdown("---")

        # Plan mode selector
        plan_mode = st.radio(
            "Plan mode:",
            ["Standard (predefined path)", "Piano di Studi Individuale (PSI)"],
            index=0,
            help="PSI includes only Curricular Exam I from the chosen sub path; you will select 3 free-choice exams.",
        )
        plan_is_psi = plan_mode.endswith("(PSI)")

        main_paths = ["Select Main Path"] + list(st.session_state.catalog.keys())
        main_choice = st.selectbox(
            "🧭 Choose Main Path:",
            main_paths,
            index=0,
            placeholder="Type to search main paths…",
            help="Start typing to quickly search.",
        )

        if main_choice != "Select Main Path":
            sub_paths = ["Select Sub Path"] + list(st.session_state.catalog[main_choice].keys())
            sub_choice = st.selectbox(
                "📂 Choose Sub Path:",
                sub_paths,
                index=0,
                placeholder="Type to search sub paths…",
                help="Start typing to quickly search.",
            )
        else:
            sub_choice = "Select Sub Path"

        if main_choice != "Select Main Path" and sub_choice != "Select Sub Path":
            st.write("### 📚 Your Curricular Courses:")
            curr_courses = st.session_state.catalog[main_choice][sub_choice]
            if plan_is_psi:
                c = curr_courses[0]
                st.markdown(f"- **Curricular 1: {c['name']}** — `{c['code']}` • **{c['cfu']} CFU** • {c['dept']} • Year: {c['year']} • Semester: {c['semester']}")
                st.info("You are in PSI mode: only Curricular Exam I is included. Select 3 free-choice exams to reach at least 60 CFU.")
            else:
                for idx, c in enumerate(curr_courses, start=1):
                    st.markdown(
                        f"- **Curricular {idx}: {c['name']}** — `{c['code']}` • **{c['cfu']} CFU** • {c['dept']} • Year: {c['year']} • Semester: {c['semester']}"
                    )

            recs = recommended_by_path.get(main_choice, {}).get(sub_choice, [])
            if recs:
                st.info("**Recommended free choice exams for this sub path:** - " + " - ".join(recs))

            n_free_required = 3 if plan_is_psi else 2
            st.write(f"### 🎯 Select {n_free_required} Free Choice Courses:")
            free_labels = [course_label(c) for c in st.session_state.free_choice_courses]
            free_choice_selection_labels = st.multiselect(
                f"Choose {n_free_required} Free Courses:",
                free_labels,
                max_selections=n_free_required,
                placeholder="Type to search free-choice courses…",
                help=f"Start typing to search; select exactly {n_free_required}.",
            )
            selected_free = [
                c for c in st.session_state.free_choice_courses if course_label(c) in free_choice_selection_labels
            ]

            fixed_total = sum(x["cfu"] for x in FIXED_COMPONENTS)
            curricular_list = [curr_courses[0]] if plan_is_psi else curr_courses
            curricular_total = sum(c["cfu"] for c in curricular_list)
            free_total = sum(c["cfu"] for c in selected_free)
            current_total = fixed_total + curricular_total + free_total

            st.caption(
                f"Planned CFUs so far: Curricular **{curricular_total}**, Free-choice **{free_total}**, Fixed components **{fixed_total}** → **{current_total}/60 CFU**"
            )

            # PSI: warn if total is below 60
            if plan_is_psi and current_total < 60:
                st.error(f"Your selections total {current_total} CFU. In PSI you must reach at least 60 CFU. Please add/change free-choice exams.")

            # Warn if total CFUs exceed 60 (e.g., due to a 12-CFU free-choice selection)
            if current_total > 60:
                st.error(f"Your selections exceed 60 CFU by {current_total - 60} CFU. Please adjust your free-choice exams or consult the coordinator.")

            can_generate = (len(selected_free) == n_free_required) and (not plan_is_psi or current_total >= 60)
            # Choose export format
            export_format = st.radio("Output format", ["PDF", "Word (.docx)"], index=0, horizontal=True)
            word_allowed = True
            if export_format == "Word (.docx)":
                word_pw = st.text_input("🔒 Password for Word export", type="password", key="wordpw")
                word_allowed = teacher_logged_in or (word_pw == "ISPD03")
            button_label = "📄 Generate PDF" if export_format == "PDF" else "📄 Generate Word (.docx)"
            if can_generate and st.button(button_label, disabled=(export_format == "Word (.docx)" and not word_allowed)):
                dob_str = dob.strftime("%d/%m/%Y") if hasattr(dob, 'strftime') else str(dob)
                if plan_is_psi:
                    ordered_courses = [
                        curr_courses[0],
                        selected_free[0],
                        selected_free[1],
                        selected_free[2],
                        FIXED_COMPONENTS[0],
                        FIXED_COMPONENTS[1],
                        FIXED_COMPONENTS[2],
                    ]
                    st.warning("This is a Piano di Studi Individuale and must be approved by the Commissione. The document will be watermarked 'To Be Approved'.")
                else:
                    ordered_courses = [
                        curr_courses[0],
                        curr_courses[1],
                        selected_free[0],
                        selected_free[1],
                        FIXED_COMPONENTS[0],
                        FIXED_COMPONENTS[1],
                        FIXED_COMPONENTS[2],
                    ]
                buf = build_study_plan_docx(
                    name=name,
                    matricula=matricula,
                    pob=pob,
                    dob_str=dob_str,
                    phone=phone,
                    email=email,
                    academic_year=academic_year,
                    year_of_degree=year_of_degree,
                    degree_type=degree_type,
                    degree_name=degree_name,
                    main_path=main_choice,
                    sub_path=(sub_choice + " — Piano di Studi Individuale" if plan_is_psi else sub_choice),
                    courses=ordered_courses,
                    watermark_text=("To Be Approved" if plan_is_psi else None),
                )
                if export_format == "PDF":
                    pdf_buf = convert_docx_to_pdf_bytes(buf)
                    if pdf_buf is None:
                        st.error("PDF conversion failed or not available. Please install `docx2pdf` and ensure Microsoft Word is available on this machine.")
                    else:
                        fname = f"Piano_di_Studi_{matricula or 'studente'}.pdf"
                        st.download_button("⬇ Download PDF", data=pdf_buf.getvalue(), file_name=fname, mime="application/pdf")
                else:
                    if not word_allowed:
                        st.error("Incorrect password for Word export.")
                    else:
                        fname = f"Piano_di_Studi_{matricula or 'studente'}.docx"
                        st.download_button("⬇ Download Word Document", data=buf, file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            elif len(free_choice_selection_labels) != n_free_required:
                st.warning(f"⚠ Please select exactly {n_free_required} free choice courses.")


if __name__ == "__main__":
    main()
